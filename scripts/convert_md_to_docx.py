"""
convert_md_to_docx.py
Chuyển đổi file Markdown sang DOCX với đầy đủ định dạng:
  - Tiêu đề (Heading 1-6) kể cả khi bọc trong **...**
  - Đoạn văn: in đậm **..**, in nghiêng *.., inline code `..`
  - Danh sách gạch đầu dòng và đánh số (hỗ trợ lồng nhau)
  - Bảng Markdown
  - Hình ảnh base64 nhúng trực tiếp
  - Lề theo chuẩn luận văn: trên 2cm, dưới 2cm, trái 3cm, phải 2cm
"""
import os
import re
import io
import base64
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Cấu hình đường dẫn
# ---------------------------------------------------------------------------
MD_PATH_VN   = "g:\\My Drive\\DoAnTotNghiep\\camera-ai\\reports\\DTC225210134_Nguy\u1ec5n Thanh Tu\u00e2n_CNTTK21CLC.docx.md"
DOCX_PATH_VN = "g:\\My Drive\\DoAnTotNghiep\\camera-ai\\reports\\DTC225210134_Nguy\u1ec5n Thanh Tu\u00e2n_CNTTK21CLC_converted.docx"


# ---------------------------------------------------------------------------
# Tiện ích chuỗi
# ---------------------------------------------------------------------------
# Loại bỏ các dấu ** __ * _ bọc ngoài một đoạn text (dùng cho heading)
_STRIP_MD_BOLD_ITALIC = re.compile(r'^(\*{1,3}|_{1,3})(.*?)(\*{1,3}|_{1,3})$')

def strip_md_bold_italic(text: str) -> str:
    """Bỏ lớp bold/italic bọc ngoài cùng của heading text."""
    text = text.strip()
    m = _STRIP_MD_BOLD_ITALIC.match(text)
    while m and m.group(1) == m.group(3):
        text = m.group(2).strip()
        m = _STRIP_MD_BOLD_ITALIC.match(text)
    return text


def clean_text(text: str) -> str:
    """
    Loại bỏ cú pháp Markdown không cần thiết để lấy plain text:
      - Link: [text](url) → text
      - Link reference: [text][ref] → text
      - Escaped chars: \\! → !, v.v.
      - Anchor header suffix: {#...} → ''
    """
    # Anchor suffix {#...}
    text = re.sub(r'\{#[^}]*\}', '', text)
    # Inline image: ![alt](url) → alt
    text = re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'\1', text)
    # Inline image ref: ![alt][ref] → alt
    text = re.sub(r'!\[([^\]]*)\]\[[^\]]*\]', r'\1', text)
    # Link: [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    # Link ref: [text][ref] → text
    text = re.sub(r'\[([^\]]+)\]\[[^\]]*\]', r'\1', text)
    # Bare link reference: [text] → text (standalone bracket)
    text = re.sub(r'\[([^\]]+)\](?!\()', r'\1', text)
    # Escaped markdown chars
    text = re.sub(r'\\([\\`*_{}\[\]()#+\-.!|>])', r'\1', text)
    return text.strip()


# ---------------------------------------------------------------------------
# Thêm text có định dạng bold/italic/code vào paragraph
# ---------------------------------------------------------------------------
# Regex tách các token định dạng inline — dùng non-greedy để tránh over-match
_INLINE_PATTERN = re.compile(
    r'(\*\*\*.+?\*\*\*'     # ***bold+italic*** (non-greedy)
    r'|\*\*(?!\*).+?(?<!\*)\*\*'  # **bold** (non-greedy, không phải ***)
    r'|__(?!_).+?(?<!_)__'       # __bold__
    r'|\*(?!\*)[^*\n]+?(?<!\*)\*' # *italic*
    r'|_(?!_)[^_\n]+?(?<!_)_'    # _italic_
    r'|`[^`\n]+`)'                # `code`
)


def add_formatted_runs(paragraph, text: str, base_size: int = 13,
                       force_bold: bool = False, force_italic: bool = False):
    """
    Phân tích inline markdown và thêm các run có định dạng vào paragraph.
    Trước tiên clean link/escape, sau đó xử lý bold/italic/code.
    """
    # Làm sạch links và escaped chars
    text = clean_text(text)
    if not text:
        return

    parts = _INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(base_size)

        if part.startswith('***') and part.endswith('***') and len(part) > 6:
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif (part.startswith('**') and part.endswith('**') and len(part) > 4):
            run.text = part[2:-2]
            run.bold = True
        elif (part.startswith('__') and part.endswith('__') and len(part) > 4):
            run.text = part[2:-2]
            run.bold = True
        elif (part.startswith('*') and part.endswith('*') and len(part) > 2
              and not part.startswith('**')):
            run.text = part[1:-1]
            run.italic = True
        elif (part.startswith('_') and part.endswith('_') and len(part) > 2
              and not part.startswith('__')):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run.text = part[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
        else:
            run.text = part

        if force_bold and not run.bold:
            run.bold = True
        if force_italic and not run.italic:
            run.italic = True


# ---------------------------------------------------------------------------
# Hỗ trợ bảng
# ---------------------------------------------------------------------------
def is_separator_row(cells):
    """Kiểm tra dòng phân cách căn chỉnh |:---|:---| trong bảng Markdown."""
    for c in cells:
        c = c.strip()
        if not re.match(r'^:?-+:?$', c):
            return False
    return True


def process_table(doc, table_lines):
    """Chuyển đổi các dòng Markdown table thành bảng Word."""
    rows = []
    for line in table_lines:
        # Tách cell
        raw = line.strip()
        if raw.startswith('|'):
            raw = raw[1:]
        if raw.endswith('|'):
            raw = raw[:-1]
        cells = [c.strip() for c in raw.split('|')]
        if is_separator_row(cells):
            continue
        rows.append([clean_text(c) for c in cells])

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    # Đảm bảo mỗi dòng đủ số cột
    for r in rows:
        while len(r) < num_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for r_idx, row_cells in enumerate(rows):
        for c_idx, val in enumerate(row_cells):
            cell = table.cell(r_idx, c_idx)
            # Xoá nội dung mặc định
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if r_idx == 0:
                run.bold = True


# ---------------------------------------------------------------------------
# Hỗ trợ ảnh
# ---------------------------------------------------------------------------
def insert_image(doc, b64_str: str, alt_text: str = ''):
    """Giải mã base64 và chèn ảnh vào tài liệu, thêm caption nếu có."""
    try:
        img_bytes = base64.b64decode(b64_str)
        img_stream = io.BytesIO(img_bytes)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(5.5))
        if alt_text and alt_text.strip() and not alt_text.lower().startswith('image'):
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(3)
            cap.paragraph_format.space_after = Pt(6)
            cr = cap.add_run(alt_text.strip())
            cr.italic = True
            cr.font.name = 'Times New Roman'
            cr.font.size = Pt(10)
    except Exception as exc:
        print(f"[WARN] Khong the chen anh: {exc}")


# ---------------------------------------------------------------------------
# Phát hiện loại dòng
# ---------------------------------------------------------------------------
# Heading: # đến ######, nội dung có thể bọc trong **...**
_HEADING_RE = re.compile(r'^(#{1,6})\s+(.*)')
# Bullet: -, *, + với tuỳ chọn indent
_BULLET_RE  = re.compile(r'^(\s*)([-*+])\s+(.*)')
# Number: 1. 2. 3. ...
_NUMBER_RE  = re.compile(r'^(\s*)(\d+)\.\s+(.*)')
# Inline image reference: ![alt][id] hoặc ![alt](url)
_IMG_REF_RE = re.compile(r'!\[([^\]]*)\]\[([^\]]+)\]')
_IMG_URL_RE = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
# Table
_TABLE_RE   = re.compile(r'^\s*\|')
# Dòng toàn bộ là dòng horizontal rule
_HR_RE      = re.compile(r'^[-*_]{3,}\s*$')


def detect_indent_level(indent_str: str) -> int:
    """Tính mức lồng nhau từ chuỗi indent (2 spaces hoặc 1 tab = 1 cấp)."""
    spaces = indent_str.replace('\t', '  ')
    return len(spaces) // 2


# ---------------------------------------------------------------------------
# Thiết lập style tài liệu
# ---------------------------------------------------------------------------
def setup_document_styles(doc):
    """Cấu hình margins, font Normal và font Heading."""
    # Margins chuẩn luận văn
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2)

    # Normal style
    style_normal = doc.styles['Normal']
    fn = style_normal.font
    fn.name = 'Times New Roman'
    fn.size = Pt(13)
    fn.color.rgb = RGBColor(0, 0, 0)
    pf = style_normal.paragraph_format
    pf.line_spacing = Pt(20)           # ~1.5 với font 13pt
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.first_line_indent = None

    # Heading styles
    heading_cfg = {
        'Heading 1': (16, True,  False),
        'Heading 2': (14, True,  False),
        'Heading 3': (13, True,  False),
        'Heading 4': (13, True,  True),
        'Heading 5': (13, False, True),
        'Heading 6': (13, False, True),
    }
    for h_name, (size, bold, italic) in heading_cfg.items():
        try:
            hs = doc.styles[h_name]
        except KeyError:
            continue
        hf = hs.font
        hf.name = 'Times New Roman'
        hf.size = Pt(size)
        hf.bold = bold
        hf.italic = italic
        hf.color.rgb = RGBColor(0, 0, 0)
        hp = hs.paragraph_format
        hp.space_before = Pt(12)
        hp.space_after  = Pt(6)
        hp.alignment    = WD_ALIGN_PARAGRAPH.LEFT
        hp.keep_with_next = True

    # List Bullet style
    try:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Times New Roman'
        lb.font.size = Pt(13)
    except KeyError:
        pass

    # List Number style
    try:
        ln = doc.styles['List Number']
        ln.font.name = 'Times New Roman'
        ln.font.size = Pt(13)
    except KeyError:
        pass


# ---------------------------------------------------------------------------
# Xử lý một dòng heading
# ---------------------------------------------------------------------------
def add_heading_line(doc, hashes: str, raw_text: str):
    """Thêm heading vào doc, tự động bỏ lớp **...** bọc ngoài."""
    level = min(len(hashes), 6)
    # Loại bỏ anchor {#...}
    raw_text = re.sub(r'\{#[^}]*\}', '', raw_text).strip()
    # Loại bỏ ** __ bọc ngoài
    raw_text = strip_md_bold_italic(raw_text)
    # Loại bỏ link markdown
    raw_text = clean_text(raw_text)
    doc.add_heading(raw_text, level=level)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    # Reconfigure stdout/stderr sang UTF-8 để tránh lỗi encode trên Windows
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')

    md_path   = MD_PATH_VN
    docx_path = DOCX_PATH_VN
    if not os.path.exists(md_path):
        print("ERROR: Khong tim thay file markdown: " + repr(md_path))
        sys.exit(1)

    print(f"Doc file: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    print(f"  Tong so dong: {len(lines)}")

    # ------------------------------------------------------------------
    # Bước 1: Tách định nghĩa ảnh base64 ra khỏi nội dung
    # ------------------------------------------------------------------
    body_lines = []
    image_defs = {}   # id -> base64_data
    # Dòng định nghĩa ảnh: [id]: <data:image/...;base64,...>
    _IMG_DEF_RE = re.compile(
        r'^\[([^\]]+)\]:\s*<?data:image/[^;]+;base64,([^>]+)>?'
    )
    for line in lines:
        stripped = line.strip()
        m = _IMG_DEF_RE.match(stripped)
        if m:
            img_id = m.group(1)
            b64    = m.group(2).strip().rstrip('>')
            image_defs[img_id] = b64
        else:
            body_lines.append(line)

    print(f"  Anh base64 trich xuat: {len(image_defs)}")
    print(f"  Dong noi dung: {len(body_lines)}")

    # ------------------------------------------------------------------
    # Bước 2: Tạo document và cấu hình style
    # ------------------------------------------------------------------
    doc = Document()
    setup_document_styles(doc)

    # ------------------------------------------------------------------
    # Bước 3: Parse và chuyển đổi từng dòng
    # ------------------------------------------------------------------
    i = 0
    total = len(body_lines)

    while i < total:
        line     = body_lines[i]
        stripped = line.rstrip('\n').rstrip('\r')
        s        = stripped.strip()

        # --- Dòng trắng ---
        if not s:
            i += 1
            continue

        # --- Horizontal rule ---
        if _HR_RE.match(s):
            i += 1
            continue

        # --- Table ---
        if _TABLE_RE.match(stripped):
            t_lines = []
            while i < total and _TABLE_RE.match(body_lines[i].rstrip('\n').rstrip('\r')):
                t_lines.append(body_lines[i].strip())
                i += 1
            process_table(doc, t_lines)
            continue

        # --- Heading ---
        hm = _HEADING_RE.match(s)
        if hm:
            add_heading_line(doc, hm.group(1), hm.group(2))
            i += 1
            continue

        # --- Bullet list ---
        bm = _BULLET_RE.match(stripped)
        if bm:
            indent = bm.group(1)
            content = bm.group(3)
            content = clean_text(content)
            level = detect_indent_level(indent)
            style = 'List Bullet' if level == 0 else 'List Bullet 2' if level == 1 else 'List Bullet 3'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(18)
            add_formatted_runs(p, content)
            i += 1
            continue

        # --- Numbered list ---
        nm = _NUMBER_RE.match(stripped)
        if nm:
            indent  = nm.group(1)
            content = nm.group(3)
            content = clean_text(content)
            level = detect_indent_level(indent)
            style = 'List Number' if level == 0 else 'List Number 2' if level == 1 else 'List Number 3'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(18)
            add_formatted_runs(p, content)
            i += 1
            continue

        # --- Ảnh inline reference: ![alt][id] ---
        img_m = _IMG_REF_RE.search(s)
        if img_m:
            alt    = img_m.group(1)
            img_id = img_m.group(2)
            if img_id in image_defs:
                insert_image(doc, image_defs[img_id], alt)
                i += 1
                continue

        # --- Ảnh inline URL: ![alt](url_or_data) ---
        img_u = _IMG_URL_RE.search(s)
        if img_u:
            alt    = img_u.group(1)
            src    = img_u.group(2)
            if src in image_defs:
                insert_image(doc, image_defs[src], alt)
                i += 1
                continue
            # Nếu src là data URI trực tiếp
            if src.startswith('data:image'):
                try:
                    _, b64part = src.split(',', 1)
                    insert_image(doc, b64part, alt)
                    i += 1
                    continue
                except Exception:
                    pass

        # --- Đoạn văn thông thường ---
        # Thu thập dòng kết thúc bằng 2 space (hard line break trong MD)
        # nhưng KHÔNG gộp nếu dòng tiếp theo bắt đầu heading/list/table/bold-block
        para_lines = [s]
        _NEXT_IS_BLOCK_RE = re.compile(
            r'^(#{1,6}\s|[-*+]\s|\d+\.\s|\|.*\||\*\*[^*])'  # heading, list, table, bold-paragraph
        )
        while stripped.endswith('  '):
            i += 1
            if i >= total:
                break
            next_line = body_lines[i].rstrip('\n').rstrip('\r')
            next_s = next_line.strip()
            if not next_s:
                break
            # Dừng nếu dòng tiếp theo là block element mới
            if (_HEADING_RE.match(next_s)
                    or _TABLE_RE.match(next_line)
                    or _BULLET_RE.match(next_line)
                    or _NUMBER_RE.match(next_line)
                    or _NEXT_IS_BLOCK_RE.match(next_s)):
                break
            para_lines.append(next_s)
            stripped = next_line

        full_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        add_formatted_runs(p, full_text)
        i += 1

    # ------------------------------------------------------------------
    # Bước 4: Lưu file
    # ------------------------------------------------------------------
    print(f"Luu tai lieu: {docx_path}")
    doc.save(docx_path)
    print("Hoan thanh! File da duoc tao thanh cong.")


if __name__ == '__main__':
    main()
